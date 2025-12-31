from openpyxl import Workbook, load_workbook
from docx import Document
print("=== Excel文件操作 ===")
wb = Workbook()
ws = wb.active
ws.title = "学生成绩"
headers = ["姓名","语文","数学","英语","总分"]
for col, header in enumerate(headers, start=1):
    ws.cell(row=1, column=col, value=header)

students = [
    ["梁仕林", 88, 92, 88, 255],
    ["李四", 90, 88, 95, 273],
    ["王五", 76, 95, 82, 253]
]

for row, student in enumerate(students, start=2):
    for col, value in enumerate(student, start=1):
        ws.cell(row=row, column=col, value=value)

excel_path = "学生成绩表.xlsx"
wb.save(excel_path)
print(f"Excel文件已创建: {excel_path}")

wb_load = load_workbook(excel_path)
ws_load = wb_load["学生成绩"]
print("读取Excel数据: ")
for row in ws_load.iter_rows(values_only=True):
    print(row)

print("\n=== Word文件操作 ===")
doc = Document()
doc.add_heading("学生成绩报告", level=1)
doc.add_paragraph("以下是2025年秋季学期学生成绩统计: ")
table = doc.add_table(rows=4, cols=5)
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header
for row_idx, student in enumerate(students, start=1):
    row_cells = table.rows[row_idx].cells
    for col_idx, value in enumerate(student):
        row_cells[col_idx].text = str(value)

doc.add_page_break()

word_path = "学生成绩报告.docx"
doc.save(word_path)
print(f"Word文件已创建: {word_path}")