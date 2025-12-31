import os
import shutil
from openpyxl import load_workbook, Workbook
from docx import Document
def read_excel_scores(excel_path):
    try:
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"输入文件'{excel_path}'不存在")
        wb = load_workbook(excel_path)
        ws = wb.active
        headers = []
        for col in range(1, 5):
            header = ws.cell(row=1, column=col).value
            if not header:
                raise KeyError("Excel表头不完整，需包含'姓名、语文、数学、英语'")
            headers.append(header.strip())
        required_headers = ["姓名","语文","数学","英语"]
        if headers != required_headers:
            raise KeyError(f"表头错误，需为{required_headers}，实际为{headers}")
        students = []
        for row in range(2, ws.max_row + 1):
            name = ws.cell(row=row, column=1).value
            if not name:
                continue
            scores = []
            for col in range(2, 5):
                score = ws.cell(row=row, column=col).value
                if not isinstance(score, (int, float)):
                    raise ValueError(f"第{row}行'{name}'的{headers[col-1]}成绩非数字：{score}")
                scores.append(score)
            chinese, math, english = scores
            total = chinese + math + english
            avg = round(total / 3, 1)
            students.append({
                "name": name,
                "chinese": chinese,
                "math": math,
                "english": english,
                "total": total,
                "avg": avg
            })
        students_sorted = sorted(students, key=lambda x: x["total"], reverse=True)
        print(f"成功读取{len(students_sorted)}名学生成绩，已按总分降序排序")
        return students_sorted
    except FileNotFoundError as e:
        print(f"文件错误: {e}")
        return None
    except KeyError as e:
        print(f"Excel格式错误: {e}")
        return None
    except ValueError as e:
        print(f"数据错误: {e}")
        return None
    except Exception as e:
        print(f"未知错误: {e}")
        return None

def export_scores(students, output_dir="成绩数据"):
    if not students:
        print("无学生数据，无法导出")
        return False
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    os.makedirs(output_dir)
    print(f"已创建目标文件夹: {output_dir}")

    txt_path = os.path.join(output_dir, "sorted_scores.txt")
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("学生成绩表（按总分降序）\n")
            f.write("=" * 50 + "\n")
            for idx, stu in enumerate(students, start=1):
                line = f"{idx}. {stu['name']}: 总分={stu['total']}, 平均分={stu['avg']}, 语文={stu['chinese']}, 数学={stu['math']}, 英语={stu['english']}\n"
                f.write(line)
        print(f"文本文件已导出: {txt_path}")
    except Exception as e:
        print(f"文本文件导出错误: {e}")
        return False

    excel_path = os.path.join(output_dir, "output_scores.xlsx")
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "排序后成绩"
        headers = ["排名","姓名","语文","数学","英语","总分","平均分"]
        for col, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col, value=header)
        for row, stu in enumerate(students, start=2):
            ws.cell(row=row, column=1, value=row-1)
            ws.cell(row=row, column=2, value=stu["name"])
            ws.cell(row=row, column=3, value=stu["chinese"])
            ws.cell(row=row, column=4, value=stu["math"])
            ws.cell(row=row, column=5, value=stu["english"])
            ws.cell(row=row, column=6, value=stu["total"])
            ws.cell(row=row, column=7, value=stu["avg"])
        wb.save(excel_path)
        print(f"Excel文件已导出: {excel_path}")
    except Exception as e:
        print(f"Excel文件导出错误: {e}")
        return False

    word_path = os.path.join(output_dir, "score_report.docx")
    try:
        doc = Document()
        doc.add_heading("学生成绩管理报告", level=1)
        total_students = len(students)
        avg_chinese = round(sum(stu["chinese"] for stu in students) / total_students, 1)
        avg_math = round(sum(stu["math"] for stu in students) / total_students, 1)
        avg_english = round(sum(stu["english"] for stu in students) / total_students, 1)
        top_stu = students[0]
        bottom_stu = students[-1]
        stats_paragraph = f"""统计信息：
1. 学生总数：{total_students}人
2. 各科平均分：语文-{avg_chinese}分，数学-{avg_math}分，英语-{avg_english}分
3. 总分最高：{top_stu['name']}（{top_stu['total']}分）
4. 总分最低：{bottom_stu['name']}（{bottom_stu['total']}分）
"""
        doc.add_paragraph(stats_paragraph)
        doc.add_heading("排序后成绩表", level=2)
        table = doc.add_table(rows=total_students + 1, cols=7)
        hdr_cells = table.rows[0].cells
        for col, header in enumerate(headers):
            hdr_cells[col].text = header
        for row_idx, stu in enumerate(students, start=1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = str(row_idx)
            row_cells[1].text = stu["name"]
            row_cells[2].text = str(stu["chinese"])
            row_cells[3].text = str(stu["math"])
            row_cells[4].text = str(stu["english"])
            row_cells[5].text = str(stu["total"])
            row_cells[6].text = str(stu["avg"])
        doc.save(word_path)
        print(f"Word报告已导出: {word_path}")
        return True
    except Exception as e:
        print(f"Word文件导出错误: {e}")
        return False

if __name__ == "__main__":
    print("=== 学生成绩管理系统 ===")
    input_excel = "input_scores.xlsx"
    print(f"\n1. 正在读取输入文件: {input_excel}")
    students_data = read_excel_scores(input_excel)
    if students_data:
        print(f"\n2. 正在导出成绩文件...")
        success = export_scores(students_data)
        if success:
            print("\n3. 所有文件导出完成！")
        else:
            print("\n3. 文件导出失败！")
    else:
        print("\n2. 读取数据失败，无法导出！")